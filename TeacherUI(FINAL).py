
import streamlit as st
import numpy as np
import pandas as pd 
import requests
from requests.auth import HTTPBasicAuth
import re
import docx
import io
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter#process_pdf
from pdfminer.pdfpage import PDFPage
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
import uuid
from io import StringIO

if 'start' not in st.session_state:
    print('session initialised...')
    st.session_state['start'] = 1
    st.session_state.updated = 0

if 'doc' not in st.session_state:
    st.session_state['doc'] = {"name": 'noname', "value": []}
    
if st.session_state.start:
    print('all_answers and final_data reset')
    all_answers = []
    final_data = ''
 
print("------***---->st.session_state.doc: ",st.session_state['doc'])

def extract_docxtext(doc):
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)

def get_pdf_searchable_pages(fname):
    # pip install pdfminer
    
    searchable_pages = []
    non_searchable_pages = []
    page_num = 0
    with open(fname, 'rb') as infile:

        for page in PDFPage.get_pages(infile):
            page_num += 1
            if 'Font' in page.resources.keys():
                searchable_pages.append(page_num)
            else:
                non_searchable_pages.append(page_num)
    if page_num > 0:
        if len(searchable_pages) == 0:
            print(f"Document '{fname}' has {page_num} page(s). "
                  f"Complete document is non-searchable")
        elif len(non_searchable_pages) == 0:
            print(f"Document '{fname}' has {page_num} page(s). "
                  f"Complete document is searchable")
        else:
            print(f"searchable_pages : {searchable_pages}")
            print(f"non_searchable_pages : {non_searchable_pages}")
    else:
        print(f"Not a valid document")

def extract_pdftext(reader):
    num_pages = len(reader.pages)
    full_text = []

    return '\n'.join(full_text)


def pdf_to_text(fp):

    # PDFMiner boilerplate
    rsrcmgr = PDFResourceManager()
    sio = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, sio, codec=codec, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    # Extract text
    #fp = file(pdfname, 'rb')
    for page in PDFPage.get_pages(fp):
        interpreter.process_page(page)
    fp.close()

    # Get text from StringIO
    text = sio.getvalue()

    # Cleanup
    device.close()
    sio.close()

    return text

def getmultitagtext(Inp_STR, tag="Question"):

    #tag = "Question"

    matches = re.findall(r"<"+tag+">(.*?)</"+tag+">", Inp_STR, flags=re.DOTALL)
    return matches

@st.dialog("Download Questions and Answers")
def create_doc(questions, answers, choices):
    # IMPORTANT: Cache the conversion to prevent computation on every rerun

    global final_data

    final_data = '''# List Of Questions and Model Answers\n\n'''
    i=0 
    for c in choices:
        if c == "Yes":
            final_data += "**Question** {}: {}".format(i+1, questions[i])
            final_data += '\n\n'
            final_data += "**Answer** {}: {}".format(i+1, answers[i])
            final_data += '\n\n'
            i += 1
    
    st.write(f"Enter a filename:")
    name = st.text_input("Example: Assessment.md")
    if st.button("Submit"):
        with open(name, "w") as f:
            f.write(final_data)
        st.session_state.doc = {"name": name, "value": final_data}
        #st.rerun()
    
    print("******** Q/A Updated *********")
    return 1
    

def download_doc(questions, answers, choices, results, doc_name):
    # IMPORTANT: Cache the conversion to prevent computation on every rerun

    global final_data

    final_data = '''# List Of Questions and Model Answers\n'''
    final_data += f"## (From {doc_name})\n\n"
    i=0 
    counter = 1
    for c in choices:
        if c == "Yes":
            final_data += "**Question** {}: {}".format(counter, questions[i])
            final_data += '\n\n'
            final_data += "**Answer** {}: {}".format(counter, answers[i])
            final_data += '\n\n'
            counter += 1
        i += 1

    suffix = str(uuid.uuid4())[:8]
    name = f'QnA_{suffix}.md'
    with open(name, "w") as f:
        f.write(final_data)
    st.session_state.doc = {"name": doc_name, "value": results}
    st.session_state.updated = 0
    
    print("******** Q/A Updated *********")
    return 1

#Heading
HTML = '''<h1 style="color:green;">Teacher Portal</h1>'''
st.markdown(HTML, unsafe_allow_html = True)

st.title("Generate Questions From Documents")
st.markdown(" **Instructions**: Select the questions that you would like to keep and include a model answer")


    
with st.sidebar:
    # Add a drop down selector for question level
    

    uploaded_files = st.file_uploader("Upload transcriptions here", accept_multiple_files=True)
    print('------------> session.start: ', st.session_state.start)
    st.session_state.start = 0
    print('uploaded file ----->', uploaded_files)

    if len(uploaded_files) > 0: 
        full_text = []
        for uploaded_file in uploaded_files:
            bytes_data = uploaded_file.read()

            if uploaded_file.name.endswith('.docx'):
            #  st.write("filename:", uploaded_file.name)
                document = docx.Document(io.BytesIO(bytes_data))
                full_text = extract_docxtext(document)
                #stringio = StringIO(bytes_data.decode("utf-8"))
            # st.write(extract_text(document))
            
            if uploaded_file.name.endswith('.pdf'):
                with io.BytesIO(bytes_data) as open_pdf_file:
                    #reader = PdfReader(open_pdf_file)
                    #full_text = extract_pdftext(reader)
                    full_text = pdf_to_text(open_pdf_file)

        option_level = st.selectbox("Type of questions?", ("MCQ", "Open-ended"))

        with st.spinner("Working on it..."):
            url = 'https://3x8htsr26h.execute-api.ap-southeast-1.amazonaws.com/test/ntc-teacher'
            headers = {'Accept': 'application/json'}
            auth = HTTPBasicAuth('kRKNFaknPj2dhCHlHLACyPQJ3x9Rptn3Aywwzq64', '1234abcd')
            data = [{"level": option_level, "content": full_text}]
            req = requests.get(url, headers=headers, auth=auth, json=data)
            print(req.json())
            temp = req.json()
            if 'answer' in temp:
                all_answers = temp['answer']
            else:
                print(temp)
                all_answers = ["<question>Why nothing returned</question>, <answer>Sorry! API times out.</answer>"]
        st.success("Done! Answers returned ....")
        st.session_state.updated = 1
        uploaded_files = []
        st.session_state['doc'] = {"name": uploaded_file.name, "value": []}


df = pd.DataFrame({'first column': ["Yes","No"]})

if len(all_answers) > 0:
    print('-----> st.session_state.doc: ', st.session_state.doc)
    if ~st.session_state.updated and len(st.session_state.doc['value']) > 0:
        results = st.session_state.doc['value']
    else:
        results = []
        for ele in all_answers:
            q = getmultitagtext(ele, "question")
            a = getmultitagtext(ele, "answer")
            results.append((q,a))

    doc_name = st.session_state.doc['name']
    #print(results)
    #Default Selection
    i=0
    questions = []
    correct_answers = []
    choice = []
    with st.form(key='Q/A'):
        for question, answer in results:
            ############ Start of Q1
            st.write("Q: {}".format(question[0]))
            questions.append(question[0])
            #print(f"Question {i}: {question}")

            opt = st.selectbox(
                label='Do you want to include this question?',
                options=df['first column'],
                key=i)
            choice.append(opt)
            print(choice)

            correct_answer = st.text_area(label="Model answer: ", value=answer[0])
            correct_answers.append(correct_answer)
            #print(f"Answer {i}: {correct_answer}")
            st.divider()
            i += 1

        st.session_state.updated = 1
        ############## End of Q1
        st.form_submit_button(
            label='Download Q/A',
            on_click=download_doc, 
            args=(questions, correct_answers, choice, results, doc_name))
            

        #file_name = "Assignment.md"
        #if st.button('Download Q/A'):
        #    create_doc(questions, correct_answers, choice)

        #st.download_button(
        #    label='Download Q/A',
        #    data=final_data,
        #    file_name = "Assignment.md",
        #    on_click=create_doc, 
        #    args=(questions, correct_answers, choice)
        #    )
        
        # Remove all answers
        #all_answers = []


    
'''

'''
              


####End of Question code


















